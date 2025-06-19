from io import BytesIO
import requests
from docx import Document
from copy import deepcopy
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Insère un paragraphe après un paragraphe donné"""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def replace_text(paragraphs, key, value):
    """Remplace les clés dans les paragraphes par la valeur correspondante"""
    value = value or ""  # Assurer que value n'est jamais None
    for p in paragraphs:
        text = ''.join(run.text for run in p.runs)
        if key in text:
            new_text = text.replace(key, value)
            for run in p.runs:
                run.text = ''
            (p.runs[0] if p.runs else p.add_run()).text = new_text

def replace_everywhere(doc, key, value):
    """Remplace les clés dans tout le document : corps, tables, en-têtes/pieds de page"""
    replace_text(doc.paragraphs, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell.paragraphs, key, value)

    for section in doc.sections:
        for container in [section.header.paragraphs, section.footer.paragraphs]:
            replace_text(container, key, value)

def list_to_string(lst):
    """Transforme une liste en texte multi-ligne sans puces, pour garder la mise en forme Word"""
    return "\n".join(item for item in lst if item)

def find_experience_block(doc):
    """Localise le bloc d'expérience dans le template Word"""
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if "[date_experience]" in p.text and start is None:
            start = i
        if "[env_item]" in p.text:
            end = i
            break
    return start, end

def duplicate_experience_blocks(doc, count):
    """Duplique le bloc d'expérience autant de fois que nécessaire"""
    start, end = find_experience_block(doc)
    if start is None or end is None:
        raise ValueError("Bloc d'expérience non trouvé dans le template Word.")

    block = doc.paragraphs[start:end + 1]
    block_elements = [deepcopy(p._element) for p in block]
    duplicated = [block]
    last = block[-1]

    for _ in range(count - 1):
        new_block = []
        for elem in block_elements:
            new_elem = deepcopy(elem)
            last._element.addnext(new_elem)
            para = Paragraph(new_elem, doc)
            new_block.append(para)
            last = para
        duplicated.append(new_block)

    return duplicated

def safe_join(items, sep=", "):
    """Joins non-empty items safely"""
    return sep.join(item for item in items if item)

def format_technologies(techno_dict):
    """Formate les technologies proprement"""
    return "\n".join(
        f"{k} : {safe_join(v)}" for k, v in techno_dict.items() if v
    )

def format_formations(formations):
    """Formate les formations proprement"""
    result = []
    for f in formations:
        titre = f.get('titre') or ''
        ecole = f.get('ecole') or ''
        lieu = f.get('lieu') or ''
        formation_str = safe_join([titre, ecole, lieu], sep=", ")
        result.append(formation_str)
    return "\n".join(result)

def mapping_cv(file):
    url = "http://35.181.31.39:8000/reformulate-cv/"

    response = requests.post(url, files={'file': file})
    if response.status_code == 200:
        cv = response.json()

        # Charge le template skiils
        doc = Document("./assets/cv_template_skiils.docx")


        # Mapping simple
        mapping = {
            "[nom_prenom]": cv.get("nom_prenom", ""),
            "[metier]": cv.get("metier", "").upper(),
            "[nb_annees]": str(cv.get("annee_experience", "")),
            "[niveau_anglais]": cv.get("anglais", "").lower(),
            "[terrain_jeu]": list_to_string(cv.get("terrain_de_jeu", [])),
            "[savoir_faire_metier]": list_to_string(cv.get("savoir_faire_metier", [])),
            "[categorie_tech ] [technologie]": format_technologies(cv.get("technologies", {})),
            "[diplome ], [ecole], [lieu]": format_formations(cv.get("formations", [])),
        }

        # Remplace les tags généraux
        for key, val in mapping.items():
            replace_everywhere(doc, key, val)

        # Gestion des expériences professionnelles
        experiences = cv.get("experiences_professionnelles", [])
        if experiences:
            blocks = duplicate_experience_blocks(doc, len(experiences))

            for exp, block in zip(experiences, blocks):
                fields = {
                    "[date_experience]": f"{exp.get('date_debut', '')} – {exp.get('date_fin', '')}",
                    "[nom_entreprise]": exp.get("entreprise", ""),
                    "[poste]": exp.get("poste", ""),
                    "[contexte]": exp.get("contexte", ""),
                    "[realisation]": list_to_string(exp.get("realisations", [])),
                    "[env_item]": exp.get("environnement", ""),
                }
                for key, val in fields.items():
                    replace_text(block, key, val)

                for para in reversed(block):
                    if para.text.strip():  # ignore les vides
                        last_real_paragraph = para
                        break
                else:
                    last_real_paragraph = block[-1]  # fallback

                # Insère l'image après le vrai dernier paragraphe utile
                new_paragraph = insert_paragraph_after(last_real_paragraph)
                new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = new_paragraph.add_run()
                run.add_picture("assets/img.png", width=Inches(5))


            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output
